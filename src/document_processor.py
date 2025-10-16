"""
Word 문서 처리 모듈
- Word 문서에서 텍스트 추출
- 이미지 추출
- 텍스트 청킹
"""

import os
from pathlib import Path
from typing import List, Dict, Tuple
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image
import io


class DocumentProcessor:
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        """
        Args:
            chunk_size: 청크 크기 (문자 수)
            chunk_overlap: 청크 간 오버랩 (문자 수)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Word 문서에서 텍스트 추출"""
        try:
            doc = Document(file_path)
            full_text = []
            
            # 본문 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # 표 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            return '\n\n'.join(full_text)
        except Exception as e:
            print(f"텍스트 추출 오류 ({file_path}): {str(e)}")
            return ""
    
    def extract_images_from_docx(self, file_path: str, output_dir: str) -> List[str]:
        """Word 문서에서 이미지 추출"""
        try:
            doc = Document(file_path)
            image_paths = []
            
            # 출력 디렉토리 생성
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            
            # 문서명 기반 폴더 생성
            doc_name = Path(file_path).stem
            doc_image_dir = Path(output_dir) / doc_name
            doc_image_dir.mkdir(exist_ok=True)
            
            # 이미지 추출
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        image = Image.open(io.BytesIO(image_data))
                        
                        # 이미지 저장
                        image_filename = f"{doc_name}_{len(image_paths)}.png"
                        image_path = doc_image_dir / image_filename
                        image.save(image_path, 'PNG')
                        
                        image_paths.append(str(image_path))
                    except Exception as e:
                        print(f"이미지 추출 오류: {str(e)}")
                        continue
            
            return image_paths
        except Exception as e:
            print(f"이미지 추출 오류 ({file_path}): {str(e)}")
            return []
    
    def chunk_text(self, text: str) -> List[str]:
        """텍스트를 청크로 분할"""
        if not text:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + self.chunk_size
            
            # 문장 경계에서 자르기
            if end < text_length:
                # 마침표, 줄바꿈 등으로 문장 경계 찾기
                for delimiter in ['\n\n', '\n', '. ', '。', '! ', '? ']:
                    pos = text.rfind(delimiter, start, end)
                    if pos != -1:
                        end = pos + len(delimiter)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # 오버랩을 고려하여 다음 시작 위치 설정
            start = end - self.chunk_overlap if end < text_length else end
        
        return chunks
    
    def process_document(self, file_path: str, image_output_dir: str) -> Dict:
        """
        Word 문서 전체 처리
        
        Returns:
            {
                'file_path': str,
                'file_name': str,
                'text': str,
                'chunks': List[str],
                'images': List[str]
            }
        """
        file_name = Path(file_path).name
        
        # 텍스트 추출
        text = self.extract_text_from_docx(file_path)
        
        # 텍스트 청킹
        chunks = self.chunk_text(text)
        
        # 이미지 추출
        images = self.extract_images_from_docx(file_path, image_output_dir)
        
        return {
            'file_path': file_path,
            'file_name': file_name,
            'text': text,
            'chunks': chunks,
            'images': images
        }
    
    def process_documents_batch(self, document_dir: str, image_output_dir: str) -> List[Dict]:
        """
        여러 Word 문서를 배치로 처리
        
        Args:
            document_dir: Word 문서가 있는 디렉토리
            image_output_dir: 이미지를 저장할 디렉토리
        
        Returns:
            List of processed document dictionaries
        """
        doc_path = Path(document_dir)
        docx_files = list(doc_path.glob('*.docx'))
        
        # 임시 파일 제외 (~$로 시작하는 파일)
        docx_files = [f for f in docx_files if not f.name.startswith('~$')]
        
        print(f"총 {len(docx_files)}개의 Word 문서를 발견했습니다.")
        
        processed_docs = []
        for file_path in docx_files:
            try:
                result = self.process_document(str(file_path), image_output_dir)
                processed_docs.append(result)
                print(f"처리 완료: {file_path.name}")
            except Exception as e:
                print(f"문서 처리 오류 ({file_path.name}): {str(e)}")
                continue
        
        return processed_docs


if __name__ == "__main__":
    # 테스트 코드
    processor = DocumentProcessor()
    
    # 단일 문서 테스트
    test_doc = "../data/documents/test.docx"
    if os.path.exists(test_doc):
        result = processor.process_document(test_doc, "../data/images")
        print(f"텍스트 청크 수: {len(result['chunks'])}")
        print(f"추출된 이미지 수: {len(result['images'])}")
